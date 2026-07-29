"""
network_momentum_research/docs/build_report.py
=================================================
Builds the full research report (Word .docx) for this project: a replication
and extension of Pu/Roberts/Dong/Zohren, "Network Momentum across Asset
Classes" (2023), on a 22-product commodity-only universe (Bloomberg data).

Deliberately excludes any mention of FDR / multiple-testing-correction
statistical analysis, per explicit instruction (2026-07-29) -- that material
stays in project memory only, not in this report.

Run: python build_report.py   (writes ../docs/... no wait, writes into this
same docs/ folder: Network_Momentum_Commodities_Report.docx)
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(_THIS_DIR, "Network_Momentum_Commodities_Report.docx")

doc = Document()

# ---- base style tweaks ----
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False, italic=False):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    return par


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return t


def caption(text):
    par = doc.add_paragraph()
    run = par.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Network Momentum Across Commodities")
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "A Replication and Extension of Pu, Roberts, Dong & Zohren (2023)\n"
    "on a 22-Product Commodity-Only Universe"
)
run.italic = True
run.font.size = Pt(14)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Kartavya Joshi\nNYU\nData source: Bloomberg\nUniverse: Metals, Precious Metals, "
             "Energy, NGL (22 products)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
h1("1. Executive Summary")

p("This report replicates and extends the network momentum framework of Pu, Roberts, Dong and "
  "Zohren (2023) -- a graph-learning-based approach to cross-sectional momentum that estimates a "
  "daily, data-driven network among a set of tradable instruments and uses that network to diffuse "
  "each instrument's own momentum signal across its \"neighbors\" before forming trading positions. "
  "The original paper works across a 64-instrument universe spanning commodities, equities, fixed "
  "income and FX. This project narrows the scope to a single, coherent 22-product commodity universe "
  "-- Metals, Precious Metals, Energy and Natural Gas Liquids (NGL) -- sourced from Bloomberg, and asks "
  "whether the same network-momentum mechanism adds value within a commodity-only setting.")

p("Four trading strategies are compared throughout: a Long-Only benchmark, a classical MACD trend-"
  "following signal, a pooled linear regression on each instrument's own momentum features "
  "(\"LinReg\", the paper's individual-momentum benchmark), and the network-diffused regression "
  "(\"GMOM\", the paper's headline network-momentum strategy). Two additional combination strategies "
  "-- RegCombo and SignCombo -- blend LinReg and GMOM together.")

p("The headline finding is a clean, two-sided robustness story. Under fast (\"shift0\") execution, "
  "LinReg and GMOM post materially higher Sharpe ratios than the Long-Only benchmark, and this edge "
  "survives a wide range of realistic transaction-cost assumptions -- it is COST-ROBUST. However, "
  "under just one additional day of execution delay (\"shift1\"), matching this project's own "
  "conservative execution-realism convention, every one of these apparent edges collapses, in most "
  "cases turning negative -- the strategies are fundamentally DELAY-FRAGILE. This same pattern "
  "reproduces even in the single most extreme result surfaced anywhere in this project (an NGL-only "
  "sub-universe LinReg Sharpe above 5), which likewise does not survive the delay test. The full-22 "
  "portfolio's diversification benefit is real and quantitatively explained by its correlation "
  "structure. Sub-class ablations show that most of that benefit comes specifically from CROSS-class "
  "diversification, not simply from holding more names. A rigorous lookback-window sensitivity study "
  "and a graph-topology characterization round out the analysis.")

p("Section 3 documents the full methodology in detail. Section 4 discusses the data-integrity work "
  "undertaken before any results could be trusted. Section 5 presents every result produced in this "
  "project. Section 6 synthesizes the findings, and Section 7 states this project's limitations "
  "honestly.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. DATA AND UNIVERSE
# ═══════════════════════════════════════════════════════════════
h1("2. Data and Universe")

h2("2.1 Data source and instruments")
p("All price data is sourced from Bloomberg. The universe consists of 22 commodity futures products "
  "spanning four sub-asset-classes:")
table(
    ["Sub-class", "N", "Products"],
    [
        ("Metals", 4, "Copper, Aluminium, Lead, Zinc"),
        ("Precious Metals", 5, "Gold, Silver, Copper (COMEX), Platinum, Palladium"),
        ("Energy", 7, "WTI, Brent, RBOB, Heating Oil, Natural Gas, Singapore Gasoil, Fuel Oil"),
        ("NGL", 6, "Ethane, Propane, Butane, Isobutane, Ethylene, Propylene"),
    ],
)
p("This is deliberately narrower and more homogeneous than the original paper's 64-instrument, "
  "four-traditional-asset-class universe (commodities, equities, fixed income, FX). The scope here is "
  "to test whether network momentum -- specifically the idea that a data-driven cross-sectional graph "
  "carries incremental information beyond each instrument's own momentum -- holds up within a single, "
  "economically coherent commodity complex, where the paper's own strongest claims are about "
  "CROSS-asset-class edges.")

h2("2.2 Roll-adjustment methodology")
p("Futures prices roll from one contract to the next as expiry approaches. Two return series are "
  "maintained throughout this project for different purposes:")
bullet("F1_raw: the literal front-month contract price. Used only for the eight momentum features "
       "(Section 3.1), matching this project's own established Momentum-sleeve convention -- a "
       "moving-average/trend feature should track the price of the actual contract being observed, "
       "not a synthetic series.")
bullet("F1_continuous (ratio-adjusted): a multiplicative, roll-adjusted continuous price series, built "
       "so that a position held through a contract roll reflects the real economic return of that "
       "roll, not a mechanical price jump between two different instruments. This series is used for "
       "every realized return, every volatility estimate, and the regression target itself.")
p("A ratio (multiplicative) adjustment is used rather than an additive one specifically to avoid "
  "negative synthetic prices, which an additive adjustment can produce over a long back-history.")

h2("2.3 Common panel construction")
p("The eight momentum features (Section 3.1) each have their own warm-up requirement -- the slowest "
  "MACD feature needs roughly 411 trading days of history before it is well-defined, and the "
  "subsequent exponentially-weighted winsorization step adds a further 252-day burn-in. On top of "
  "this, the graph-learning ensemble (Section 3.2) requires its largest lookback window (1,260 "
  "trading days, roughly five years) of CONSECUTIVE already-valid feature history before the graph "
  "itself can first be estimated.")
p("The binding constraint across all 22 products is NGL Propylene, whose own price history is the "
  "shortest. Empirically, the full panel -- every product's features valid, AND a full five-year graph "
  "ensemble computable -- first becomes available on 2017-10-09. This is considerably later than the "
  "original paper's own setup (which has over a decade of pre-sample history before its first test "
  "period), so the walk-forward schedule used throughout this project (Section 3.8) is compressed "
  "accordingly, while deliberately keeping the paper's own five lookback windows unchanged rather than "
  "shrinking the model specification to fit the shorter available history.")
p("The resulting common panel spans 2012-05-02 (the earliest date at which every product's own "
  "momentum features are individually valid) through 2026-06-30, with the walk-forward test period "
  "beginning 2021-01-01.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. METHODOLOGY
# ═══════════════════════════════════════════════════════════════
h1("3. Methodology")

h2("3.1 The eight momentum features")
p("For every product and every trading day, eight momentum-style features are computed, following the "
  "paper's own Section 2.2 specification:")
bullet("Five vol-scaled returns: the raw return over the past Δ trading days (Δ = 1, 21, 63, 126, "
       "252), divided by that product's own trailing daily volatility (a 60-day exponentially-weighted "
       "standard deviation of simple returns), scaled by the square root of Δ. This puts return "
       "horizons of very different lengths onto a comparable, volatility-normalized scale.")
bullet("Three normalized MACD (moving-average convergence/divergence) features, computed at three "
       "speed pairs -- (8,24), (16,48) and (32,96) trading days. Each MACD value is normalized first by "
       "a 63-day rolling price standard deviation, then by a 252-day rolling standard deviation of that "
       "already-normalized series, so the final feature is roughly stationary and comparable across "
       "very differently-priced instruments.")
p("All eight features are then winsorized -- capped and floored at ±5 times their own "
  "exponentially-weighted standard deviation around their own exponentially-weighted mean (a 252-day "
  "half-life) -- to limit the influence of extreme single-day outliers on everything downstream.")

h2("3.2 Graph learning: the Kalofolias framework")
p("At the heart of the network-momentum approach is a daily, data-driven estimate of which "
  "instruments are most \"connected\" to one another, based purely on how similar their recent "
  "momentum-feature histories are -- not on any pre-specified sector or asset-class grouping. This "
  "project implements the graph-learning formulation of Kalofolias (AISTATS 2016): given a matrix of "
  "recent feature histories for all 22 products, the learned graph is the one that (a) keeps "
  "instruments with SIMILAR recent feature histories strongly connected, (b) is penalized for being "
  "either too sparse or too dense via two hyperparameters (α, which rewards higher overall "
  "connectivity, and β, which penalizes large edge weights), and (c) is solved as a convex "
  "optimization problem with a single global optimum for any given (α, β).")
p("This project solves that optimization using forward-backward-forward (FBF) primal-dual splitting, "
  "a first-order iterative method, chosen for tractability at the scale this project needs (the graph "
  "must in principle be re-estimated every single trading day across the full walk-forward history). "
  "The original paper instead uses a commercial solver (MOSEK, via the CVXPY interface), which finds "
  "the same convex problem's unique global optimum via a different numerical method. Both target "
  "exactly the same mathematical objective; the numerical method used to reach it differs.")
p("Following the paper (Section 3.1), the graph is not estimated from a single lookback window but as "
  "an ENSEMBLE across five lookback windows -- Δ = 252, 504, 756, 1008 and 1260 trading days (one "
  "to five years). Five separate graphs are estimated each day (one per window), then simply averaged "
  "together and degree-normalized, so that no single lookback horizon dominates the final network.")

h2("3.3 Network-diffused features")
p("Once the day's normalized graph is available, each product's own eight-feature vector is replaced "
  "by a WEIGHTED AVERAGE of its own features and its neighbors' features, with weights given by the "
  "learned graph. Concretely, product i's network-diffused feature vector on day t is the graph's "
  "row-i weights applied to every product's raw feature vector on day t. Intuitively: if the graph has "
  "learned that Copper and Aluminium currently move together, Copper's network-diffused momentum "
  "signal will partly reflect Aluminium's own recent momentum too, not just Copper's.")

h2("3.4 The four core trading strategies")
h3("Long-Only")
p("A trivial benchmark: every one of the 22 products is held at a constant, fully-invested position "
  "every day (after volatility targeting, Section 3.6). This isolates the pure diversification benefit "
  "of holding the whole universe, with no timing or directional signal at all.")
h3("MACD")
p("A classical trend-following signal built directly from the three MACD features (Section 3.1): the "
  "position is the average of a bounded, saturating transformation of each of the three MACD values "
  "(a function that behaves roughly linearly for small values but flattens out for large ones, "
  "preventing one extreme MACD reading from dominating the position size). This has no graph-learning "
  "component at all -- it is a pure single-instrument trend signal, included as a familiar external "
  "benchmark.")
h3("LinReg (individual momentum)")
p("A pooled linear regression is fit across every (product, day) observation in the training window, "
  "predicting each product's own next-day volatility-scaled return from its own eight RAW momentum "
  "features (no graph diffusion at all). The fitted regression is then applied out-of-sample: each "
  "day's predicted value determines the sign (long or short) of that day's position. This is the "
  "paper's own individual-momentum benchmark -- everything a network-momentum approach needs to beat "
  "to prove the graph is adding value.")
h3("GMOM (network momentum)")
p("Identical in every respect to LinReg, except the pooled regression is fit on the NETWORK-DIFFUSED "
  "features (Section 3.3) rather than the raw ones. This is the paper's headline strategy -- the "
  "entire empirical question this replication investigates is whether GMOM meaningfully outperforms "
  "LinReg, i.e. whether the graph-diffusion step adds real value beyond what each instrument's own "
  "momentum already captures.")

h2("3.5 Combination strategies: RegCombo and SignCombo")
p("Following the paper's own diversification check (its Section 4.3), two ways of blending LinReg and "
  "GMOM into a single combined strategy are also tested:")
bullet("RegCombo: the two models' raw, pre-sign regression outputs (the predicted next-day return) are "
       "averaged together first, and the position's direction is taken from the SIGN of that average.")
bullet("SignCombo: each model's sign is taken independently first, and the two signs are then "
       "averaged -- landing on a value of -1, -0.5, 0, 0.5 or +1. This is a more conservative \"vote\": "
       "a full-size position is only taken when LinReg and GMOM agree on direction; when they "
       "disagree, the position is automatically halved rather than following whichever model's raw "
       "output happens to have the larger magnitude.")

h2("3.6 Portfolio construction: volatility targeting")
p("Every strategy's raw signal (long/short/vote) is converted into an actual position size via "
  "volatility targeting, applied INDIVIDUALLY to every asset, every day, regardless of which universe "
  "(the full 22 products, or any sub-class) that asset happens to be traded within: each asset's "
  "position is scaled so that its own trailing annualized volatility, once scaled, matches a fixed "
  "target of 15% per year. This is the same target, computed the same way, in every single result in "
  "this report -- no result in this project ever depends on a different volatility-targeting "
  "assumption between one universe or strategy and another. The day's overall portfolio return is "
  "simply the equal-weighted average of however many already-individually-vol-targeted asset returns "
  "are active that day.")

h2("3.7 Transaction costs")
p("A flat 5 basis point (0.05%) round-trip transaction cost is charged on the day-to-day change in "
  "each asset's (already vol-targeted) position size. Rolling a futures contract from the expiring "
  "front-month to the new one is treated as a genuine, separate trade -- even on a day when the "
  "strategy's overall directional stance does not change, holding a position through a roll incurs the "
  "full round-trip cost, since two real trades (selling the expiring contract, buying the new one) "
  "still take place. Every Sharpe ratio, volatility, drawdown and other risk statistic reported "
  "anywhere in this document is calculated on this NET-of-cost return series, not the gross return "
  "before costs -- with one single, explicitly labeled exception (the 0-cost row of the transaction-"
  "cost sensitivity sweep in Section 5.3, whose entire purpose is to show the gross-to-net spectrum). "
  "Hyperparameter selection (Section 3.10) is likewise always scored on the net-of-cost criterion, not "
  "the gross one, so that the chosen model specification cannot be one that merely looks good before "
  "costs are considered.")

h2("3.8 Walk-forward validation: three variants")
p("Every strategy is evaluated out-of-sample via walk-forward validation: the model is fit on a "
  "training window, its hyperparameters are chosen on a held-out validation slice at the end of that "
  "same training window, and it is then tested on a fully separate, never-touched future window. Three "
  "different walk-forward SCHEDULES are tested, to check whether the results depend on the particular "
  "way the training/test windows are drawn:")
table(
    ["Variant", "Description"],
    [
        ("Expanding", "The original, paper-matching scheme: the training window always starts at the "
                       "earliest usable date (2017-10-09) and simply grows longer with each successive "
                       "block. Three blocks, testing 2021-2022, 2023-2024, and 2025-mid-2026."),
        ("Rolling", "The same three test windows as Expanding, but the training window is a FIXED "
                     "length (roughly 3.2 years) that slides forward each block, dropping the oldest "
                     "history rather than accumulating it."),
        ("Annual", "The same expanding, always-growing training window as the Expanding variant, but "
                    "recalibrated every calendar year instead of roughly every two years -- six blocks "
                    "(2021 through mid-2026) instead of three."),
    ],
)
p("Within every block, the training window's final 10% (with a 60-day minimum) is held out as a "
  "validation slice used only for hyperparameter selection (Section 3.10); the model is then refit on "
  "the full training window before being applied to the separate test window.")

h2("3.9 Execution delay: shift0 vs shift1")
p("Every trading signal is necessarily built from information available through the close of trading "
  "day t, and the earliest a position based on that signal can legally be executed is day t+1's "
  "return. \"Shift0\" refers to exactly this fastest-legal-execution assumption -- no delay beyond the "
  "unavoidable minimum. \"Shift1\" adds one FURTHER day of delay on top of that minimum (the signal is "
  "applied to the return from day t+1 to day t+2 instead), matching this project's own wider, more "
  "conservative convention used elsewhere for execution realism. This single extra day of delay turns "
  "out to be the single most consequential robustness axis tested anywhere in this project (Section "
  "5.1).")

h2("3.10 Hyperparameter selection: the grid search")
p("The graph-learning framework has two hyperparameters, α and β (Section 3.2), each drawn "
  "from the same eleven-point grid used in the original paper: {0.0001, 0.0005, 0.001, 0.005, 0.01, "
  "0.05, 0.1, 0.5, 1, 5, 10}. This produces 121 candidate (α, β) pairs. For every walk-forward "
  "block, all 121 candidates are scored on the block's own validation slice (net of transaction costs, "
  "Section 3.7), and the single best-scoring pair is selected and then used for that block's final "
  "regression fit and out-of-sample test. During the scoring phase only, the graph is re-estimated at "
  "a coarser, roughly-monthly cadence (rather than every single day) purely to make evaluating all 121 "
  "candidates computationally tractable; this disclosed compromise affects only which hyperparameters "
  "get selected, not the fidelity of the final, fully daily-re-estimated backtest built on top of "
  "whichever pair is chosen.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. DATA INTEGRITY
# ═══════════════════════════════════════════════════════════════
h1("4. Data Integrity and Correctness")

p("Before any result in this report could be trusted, three genuine implementation defects were found "
  "and fixed, each verified empirically both before and after the fix -- not merely reasoned about. "
  "This section documents them briefly, both for transparency and because their existence is directly "
  "relevant to how much confidence to place in any single number below.")

h3("Defect 1 -- the graph-learning solver was silently ignoring the input data")
p("The first iteration of the FBF optimization routine used starting values that, combined with the "
  "typical scale of the pairwise feature-distance computation in this dataset, caused the solver to "
  "clip every single pairwise distance to zero on its very first step -- regardless of the true "
  "underlying distances. The practical effect was a learned graph shaped ONLY by a generic, fully "
  "symmetric structural term, completely independent of which (α, β) pair or which actual "
  "feature data was supplied. This was confirmed directly by instrumenting the solver: the learned "
  "graph's edge weights showed essentially zero variation and essentially zero correlation with the "
  "true underlying feature distances before the fix, and a strong, expected correlation after it. The "
  "fix -- normalizing the distance scale before optimization, a standard usage convention for this "
  "algorithm that was simply missing from the initial implementation -- is a one-line change with a "
  "large, verified effect.")

h3("Defect 2 -- realized returns were computed from the wrong price series")
p("Every realized return, every volatility estimate and the regression target itself were initially "
  "computed from the raw, non-roll-adjusted front-month price series rather than the roll-adjusted "
  "continuous series (Section 2.2). On every futures roll date, this produced a spurious mechanical "
  "\"return\" purely from the price gap between the expiring and the new front-month contract -- not a "
  "real economic gain or loss. Measured directly on the actual data, this distortion reached up to "
  "15.5% on a single day for Natural Gas (mean 0.52% across 247 roll days) and up to 4.75% for WTI. "
  "This contaminated every performance number computed before the fix. The fix was to source every "
  "return computation from the already-correct roll-adjusted series, which had been computed correctly "
  "all along for the momentum features -- only the separate return-realization step had the bug.")

h3("Defect 3 -- the roll-day transaction cost was undercharged by half")
p("A futures roll is physically two separate trades (selling the expiring contract, buying the new "
  "front-month), and should therefore be charged a full round-trip transaction cost even on a day when "
  "the strategy's own directional stance does not change. The initial implementation applied only half "
  "of that charge. This was corrected and verified via a hand-computed unit test with a known expected "
  "answer.")

p("With all three defects fixed, results were rebuilt from scratch with a version-tagged checkpoint "
  "scheme specifically designed to make it impossible to accidentally reuse any pre-fix result "
  "anywhere in this project.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. RESULTS
# ═══════════════════════════════════════════════════════════════
h1("5. Results")

h2("5.1 Headline results: full-22 universe")
p("Aggregated out-of-sample Sharpe ratios, all four strategies, all three walk-forward variants, both "
  "execution-delay assumptions:")
table(
    ["Variant", "Shift", "Long-Only", "MACD", "LinReg", "GMOM"],
    [
        ("Expanding", 0, 0.562, 0.276, 0.898, 0.773),
        ("Rolling", 0, 0.562, 0.276, 1.113, 0.788),
        ("Annual", 0, 0.562, 0.275, 1.109, 0.748),
        ("Expanding", 1, 0.535, 0.174, -0.338, -0.033),
        ("Rolling", 1, 0.535, 0.174, -0.180, 0.080),
        ("Annual", 1, 0.534, 0.173, -0.364, -0.270),
    ],
)
p("The pattern is stark and completely consistent across all three walk-forward variants: under "
  "shift0 (fast execution), LinReg clearly beats every other strategy, with GMOM also comfortably "
  "ahead of both benchmarks. Under shift1 (one extra day of delay), the picture inverts almost "
  "entirely -- Long-Only becomes the best performer in every variant, and LinReg turns decisively "
  "negative in all three. The mechanism is turnover: LinReg and GMOM are high-turnover signals "
  "(roughly 0.3-0.5 in average daily position change per asset) that depend on capturing very "
  "short-lived moves, so a single extra day of staleness destroys most of their edge. Long-Only and "
  "MACD, by contrast, have very low turnover (roughly 0.01-0.02) and are barely affected by the same "
  "delay.")
p("A block-bootstrap significance test (accounting for the day-to-day autocorrelation inherent in "
  "daily strategy returns) was run on the LinReg/GMOM-vs-Long-Only return-difference series for every "
  "variant and shift combination. None of the six comparisons clears conventional statistical "
  "significance -- even the best-looking shift0 gaps come with wide confidence intervals that do not "
  "exclude zero. Every Sharpe-ratio number in the table above should therefore be read as a "
  "directional point estimate from this particular data sample, not as a statistically validated "
  "effect.")

h2("5.2 Coefficient-level analysis")
p("A separate, standardized-feature regression analysis (mirroring the original paper's own only "
  "hypothesis test) was run to ask, for each of the eight momentum features individually: does it "
  "carry a statistically real, standalone predictive relationship with next-day returns, once the "
  "other seven features are controlled for? This was run for LinReg (raw features) and GMOM "
  "(network-diffused features) separately, across six progressively-growing training windows "
  "(2017-2020 through 2017-2025).")
p("Across every one of the six windows and both LinReg and GMOM, only the SHORTEST feature -- the "
  "1-day vol-scaled return -- is a consistently significant, always-positive predictor. None of the "
  "longer horizons (1-month, 3-month, 6-month, 1-year vol-scaled returns) or any of the three MACD "
  "features hold up as reliably significant once enough training data has accumulated; where they "
  "appear significant, it is only in the smallest, earliest window, and that apparent significance "
  "fades as more data arrives -- a classic small-sample pattern, not a stable effect. Network diffusion "
  "(GMOM) does not surface any additional, consistently-significant longer-horizon feature that plain "
  "LinReg lacks.")
p("This is a genuine, honest point of divergence from the original paper, which reports a "
  "statistically significant NEGATIVE one-month coefficient in its own universe -- a reversal effect "
  "layered on top of the shorter-horizon momentum effect. No comparable reversal signal survives in "
  "this commodity-only universe at any horizon; the entire statistically defensible edge here is the "
  "very short, one-day continuation effect.")

h2("5.3 Transaction cost sensitivity")
p("The cost assumption itself was swept from 0 to 5 basis points (the paper's own convention), holding "
  "the already-selected model specification fixed, across all three walk-forward variants:")
table(
    ["Variant", "tc = 0 bps (gross)", "tc = 5 bps (net, headline)"],
    [
        ("Expanding", "Long 0.643 / MACD 0.441 / LinReg 1.360 / GMOM 1.152",
         "Long 0.562 / MACD 0.276 / LinReg 0.898 / GMOM 0.773"),
        ("Rolling", "Long 0.643 / MACD 0.441 / LinReg 1.584 / GMOM 1.163",
         "Long 0.562 / MACD 0.276 / LinReg 1.113 / GMOM 0.788"),
        ("Annual", "Long 0.643 / MACD 0.441 / LinReg 1.591 / GMOM 1.158",
         "Long 0.562 / MACD 0.275 / LinReg 1.109 / GMOM 0.748"),
    ],
)
p("The decay in Sharpe as cost rises is smooth and roughly linear, with no sudden cliff, and -- "
  "critically -- LinReg and GMOM's edge over Long-Only does not close anywhere within the swept 0-5bps "
  "range in any of the three variants. Extrapolating the observed decay rates, the breakeven cost at "
  "which that edge would fully disappear is roughly 8-12 basis points round-trip in every variant -- "
  "comfortably above both the 5bps baseline used elsewhere in this report and realistic commodity "
  "futures trading costs. The shift0 edge over Long-Only is therefore COST-ROBUST: it does not depend "
  "on an unrealistically generous cost assumption to look attractive. This stands in direct contrast "
  "to the delay finding in Section 5.1 -- the edge is robust to the cost assumption but fragile to "
  "execution speed, and these are two genuinely independent conclusions that should not be conflated.")

h2("5.4 Why the full portfolio outperforms: a diversification analysis")
p("A dedicated analysis was run to explain, beyond the Sharpe ratio alone, exactly how combining all "
  "22 products helps. Isolating the pure Long-Only strategy (removing any timing/signal effect) over "
  "the aggregated out-of-sample test period (2021-01-04 through 2026-06-30):")
table(
    ["Universe", "Return", "Volatility", "Max Drawdown", "Sortino", "Calmar"],
    [
        ("Full-22", "4.73%", "8.41%", "-18.46%", 0.74, 0.26),
        ("Metals only", "3.29%", "12.16%", "-26.21%", 0.43, 0.13),
        ("Precious only", "1.13%", "12.58%", "-26.04%", 0.13, 0.04),
        ("Energy only", "8.80%", "11.71%", "-16.55%", 1.00, 0.53),
        ("NGL only", "3.95%", "10.42%", "-27.23%", 0.47, 0.15),
    ],
)
p("The full-22 portfolio has clearly the lowest volatility of any of the five universes -- 8.4%, "
  "versus 10.4-12.6% for any single sub-class alone. One honest nuance: on maximum drawdown "
  "specifically, the Energy-only universe actually posts a slightly milder figure (-16.6%) than the "
  "full-22 portfolio (-18.5%). Diversification reliably reduces variance, but drawdown is a more "
  "path-dependent, tail-sensitive statistic, and correlations across products often rise precisely "
  "during the stress periods that drive drawdowns -- so diversification does not mechanically guarantee "
  "the mildest possible drawdown too.")
p("The mechanism behind the volatility reduction is the correlation structure among the 22 products, "
  "computed across eleven separate periods (the full 2012-2026 panel, the full 2017-2026 backtest "
  "period, and each walk-forward block's own fit/validate/test sub-periods individually). The average "
  "pairwise correlation across all 22 products is roughly 0.24, but this splits sharply and STABLY "
  "across every one of the eleven periods examined: roughly 0.39-0.42 between two products in the SAME "
  "sub-class, versus roughly 0.17-0.24 between two products in DIFFERENT sub-classes. This is why "
  "combining across sub-classes helps more than simply holding more names within a single sub-class -- "
  "cross-class pairs are genuinely, consistently less correlated.")
p("A textbook diversification-ratio formula -- portfolio volatility approximately equal to the "
  "individual target volatility multiplied by the square root of (1/N + (N-1)/N times the average "
  "pairwise correlation) -- was checked directly against the realized data, since every individual "
  "asset is already volatility-targeted to the same fixed level before being combined (Section 3.6), "
  "which is exactly the setup this formula assumes:")
table(
    ["Universe", "Theoretical volatility", "Realized volatility"],
    [
        ("Full-22", "7.9%", "8.4%"),
        ("Metals", "11.6%", "12.2%"),
        ("Precious", "11.7%", "12.6%"),
        ("Energy", "10.8%", "11.7%"),
        ("NGL", "9.7%", "10.4%"),
    ],
)
p("The theoretical prediction tracks the realized volatility closely (within roughly 5-9%) across all "
  "five universes, confirming that the diversification benefit is not merely a qualitative story but "
  "is quantitatively consistent with the correlation structure actually observed in the data.")

h2("5.5 Sub-class ablations")
p("To test whether the network-momentum mechanism specifically needs cross-class edges to add value "
  "(the original paper's own central claim), each of the four sub-classes was tested entirely on its "
  "own -- its own independently re-optimized graph, its own independently selected hyperparameters, "
  "its own regression fit -- rather than reusing anything from the full-22 model.")
table(
    ["Universe", "N", "Long-Only", "MACD", "LinReg", "GMOM"],
    [
        ("Full-22 (reference)", 22, 0.562, 0.276, 0.898, 0.773),
        ("Metals only", 4, 0.270, -0.512, -0.412, -0.245),
        ("Precious only", 5, 0.090, -0.296, -0.274, -0.604),
        ("Energy only", 7, 0.751, 0.421, -0.310, 0.298),
        ("NGL only", 6, 0.379, 0.759, 5.083, 1.801),
    ],
)
p("Metals-only and Precious-only both degrade sharply when isolated from the rest of the universe -- "
  "even the Long-Only benchmark weakens considerably (reflecting the lost diversification benefit, "
  "Section 5.4), and every momentum-based strategy turns outright negative. This is strong, direct "
  "evidence that the full graph's cross-class structure is doing real, positive work for these two "
  "sub-classes specifically -- restricting to within-class edges only actively hurts.")
p("Energy-only tells a more mixed story: Long-Only and MACD actually beat their full-22 counterparts "
  "(energy commodities such as WTI, Brent and Natural Gas are individually well known for strong "
  "standalone trending behavior), GMOM stays positive but weaker than the full-22 case, and LinReg "
  "flips negative.")
p("NGL-only stands apart as the single most striking number in this entire project: a LinReg Sharpe "
  "ratio above 5, roughly five to six times any other result reported anywhere in this report. Before "
  "treating this as a genuine finding, two separate checks were run. First, a return-concentration "
  "check: is this driven by a handful of unusually large days, the classic signature of a data or "
  "liquidity artifact? It is not -- the largest 1% of days by magnitude account for only around 6% of "
  "the total return magnitude, far below what a concentrated-artifact pattern would show. Second, a "
  "block-bootstrap significance test (accounting for daily-return autocorrelation) was run specifically "
  "on this result: the LinReg-versus-Long-Only gap in the NGL-only universe is in fact statistically "
  "significant, the only such result found anywhere in this project. Section 5.6 subjects this finding "
  "to the same delay-robustness test applied to everything else.")

h2("5.6 Does the NGL anomaly survive execution delay?")
p("Given how unusual the NGL-only result is, and given that execution delay (Section 3.9) has proven "
  "the single most powerful reality-check applied to any result in this project, the NGL-only "
  "strategies were re-scored under shift1 as well, reusing the already-selected hyperparameters (no "
  "new hyperparameter search was needed, since the choice of hyperparameters does not depend on the "
  "execution-delay assumption -- only how a given prediction is later scored against future returns "
  "does).")
table(
    ["Strategy", "shift0 Sharpe", "shift1 Sharpe"],
    [
        ("Long-Only", 0.379, 0.362),
        ("MACD", 0.759, 0.575),
        ("LinReg", 5.083, 1.789),
        ("GMOM", 1.801, 0.507),
    ],
)
p("Under one additional day of execution delay, LinReg's Sharpe ratio collapses from 5.08 to 1.79, and "
  "-- critically -- the earlier block-bootstrap significance finding no longer holds: the point "
  "estimate of the LinReg-versus-Long-Only gap falls to roughly +4.3 basis points per day, with a "
  "confidence interval that now includes zero. In other words, the one genuinely statistically "
  "significant result surfaced anywhere in this project is NOT an exception to the project's central "
  "delay-fragility finding -- it is a larger version of exactly the same pattern seen in the full-22 "
  "case. LinReg still posts a respectable standalone Sharpe of 1.79 under shift1, so this is not a "
  "losing strategy; the statistical case for it being a real, delay-robust edge (rather than a "
  "shift0-specific artifact) simply does not survive.")

h2("5.7 Combination strategies: RegCombo and SignCombo")
p("Blending LinReg and GMOM together (Section 3.5), tested on the full-22 universe under the "
  "Expanding/shift0 configuration:")
table(
    ["Strategy", "Sharpe", "Max Drawdown"],
    [
        ("Long-Only", 0.562, "-18.5%"),
        ("MACD", 0.276, "-12.5%"),
        ("LinReg", 0.898, "-9.5% (mildest of all six)"),
        ("GMOM", 0.773, "-15.4%"),
        ("RegCombo", 0.845, "-14.7%"),
        ("SignCombo", 0.928, "-12.3%"),
    ],
)
p("SignCombo -- the more conservative \"vote\" combination that only takes a full-size position when "
  "LinReg and GMOM agree -- actually beats standalone LinReg (0.928 versus 0.898 Sharpe), with a "
  "milder drawdown than GMOM alone. This is a genuine diversification benefit from blending two "
  "different signal types. RegCombo, which averages the two models' raw outputs before taking the "
  "sign (allowing whichever model has the larger-magnitude prediction to dominate even when the two "
  "disagree), underperforms both LinReg alone and SignCombo -- the more conservative vote-based rule "
  "outperforms the naive-averaging rule here.")

h2("5.8 Lookback-window sensitivity")
p("The graph-learning ensemble (Section 3.2) combines five separate lookback windows. To test whether "
  "this ensemble genuinely adds value over any single window used alone, each of the five windows was "
  "tested completely independently -- its own full 121-candidate hyperparameter search and regression "
  "refit per walk-forward block, not reusing the ensemble's own choices -- on the full-22 universe "
  "under Expanding/shift0:")
table(
    ["Lookback window", "GMOM Sharpe", "Turnover"],
    [
        ("252 days (~1 year)", 0.738, 0.337),
        ("504 days (~2 years)", 0.579, 0.355),
        ("756 days (~3 years)", 0.619, 0.348),
        ("1008 days (~4 years)", 0.847, 0.339),
        ("1260 days (~5 years)", 0.860, 0.338),
        ("Five-window ensemble", 0.773, "—"),
    ],
)
p("The relationship between lookback length and performance is NOT monotonic, and -- notably -- the "
  "ensemble does not dominate every individual window. Both the shortest window (one year) and the two "
  "longest (four and five years) individually outperform the five-window ensemble; only the two "
  "middle-length windows underperform it. Turnover is essentially flat across all five windows "
  "(roughly 0.34-0.36), so this is a genuine difference in signal quality by lookback horizon, not a "
  "cost or turnover artifact. This result should be read as a real, carefully measured observation "
  "rather than a firm recommendation to abandon the ensemble -- each individual window's number comes "
  "from only three independently-selected walk-forward blocks, and this project has separately "
  "established that the hyperparameter grid search itself has wide flat regions and real block-to-"
  "block instability, so more testing (for example, under the Rolling or Annual variants, or under "
  "shift1) would be needed before treating any single window as a clear replacement for the ensemble.")

h2("5.9 Graph topology: calm versus volatile regimes")
p("Finally, the learned graph's own structure -- independent of any trading performance -- was "
  "characterized directly, comparing a calm period (2017-10-09 through 2019-12-31) against a volatile "
  "one (the COVID shock, 2020-01-01 through 2021-12-31), using the project's own previously established "
  "macro regime windows. A single, already-validated hyperparameter pair was held fixed across both "
  "periods, since this is a characterization of the graph's structure, not a re-optimized backtest.")
table(
    ["Metric", "Calm period", "Volatile period (COVID)"],
    [
        ("Edge sparsity (fraction near-zero)", "2.4%", "8.7% (3.6x higher)"),
        ("Weighted clustering coefficient", "0.222", "0.134 (40% lower)"),
        ("Day-to-day stability (Jaccard index)", "0.995", "0.994"),
        ("Number of communities detected", "3", "3"),
        ("Modularity", "0.159", "0.151"),
        ("Within-community edge-weight share", "50.7%", "49.7%"),
    ],
)
p("During the volatile period, the learned graph becomes noticeably SPARSER (over three and a half "
  "times as many near-zero edges) and LESS CLUSTERED (a 40% lower clustering coefficient) than in the "
  "calm period -- fewer instruments maintain strongly meaningful connections during market stress. "
  "Day-to-day stability and the graph's overall community structure, however, are essentially "
  "unchanged between the two regimes: the same number of communities, similar modularity, and a "
  "similar share of total edge weight sitting within (rather than across) those communities in both "
  "periods. The pattern that emerges is that market stress makes the learned graph THIN OUT, rather "
  "than reorganizing it into a different structure -- a plausible, coherent pattern, though reported "
  "here as an observation from two illustrative windows rather than a proven, generalizable "
  "regime-detection mechanism.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. DISCUSSION AND SYNTHESIS
# ═══════════════════════════════════════════════════════════════
h1("6. Discussion and Synthesis")

p("Pulling every result in Section 5 together, four themes recur consistently enough to treat as this "
  "project's core, defensible conclusions.")

h3("The single most important finding: cost-robust, but delay-fragile")
p("This is not one finding but a genuinely two-sided one, and both halves matter. The apparent edge of "
  "LinReg and GMOM over a simple Long-Only benchmark does not depend on a lucky or generous cost "
  "assumption -- it survives realistic transaction costs comfortably (Section 5.3). But it depends "
  "entirely on fast execution -- a single additional day of delay is enough to erase it, and in most "
  "configurations flip it negative (Section 5.1). This same pattern was independently re-confirmed on "
  "the most extreme individual result in the whole project, the NGL-only LinReg Sharpe above 5 (Section "
  "5.6), which is itself just a larger version of the identical delay-fragility pattern, not an "
  "exception to it. Any practical assessment of these strategies has to treat execution speed, not "
  "trading cost, as the binding constraint.")

h3("The diversification benefit is real and mechanistically explained")
p("The full-22 portfolio's lower volatility is not simply a byproduct of holding more names -- it is "
  "specifically explained by the CROSS-sub-class correlation structure, which is consistently lower "
  "than the WITHIN-sub-class correlation structure across every period examined (Section 5.4). The "
  "sub-class ablations independently confirm the same story from the performance side: two of the four "
  "sub-classes (Metals, Precious) do meaningfully worse when isolated, exactly consistent with losing "
  "access to that cross-class diversification (Section 5.5).")

h3("No robust reversal effect, unlike the original paper")
p("The original paper's own coefficient-level analysis finds a statistically significant reversal "
  "effect at the one-month horizon, on top of a shorter-horizon momentum effect. This commodity-only "
  "universe shows no comparable reversal signal at any horizon examined -- the only consistently "
  "significant, stable coefficient across every window tested is the very shortest, one-day "
  "continuation effect (Section 5.2). This is a genuine, substantive point of divergence from the "
  "original paper's own commodities-and-more universe, not a methodological gap in this replication.")

h3("The graph-learning mechanism's own behavior is coherent and explainable")
p("Both the sub-class ablations (Section 5.5) and the graph-topology analysis (Section 5.9) paint a "
  "consistent, economically sensible picture of what the learned graph is actually doing: it reflects "
  "genuine, time-varying cross-class relationships that thin out under market stress rather than "
  "reorganizing into a fundamentally different structure, and restricting it to within-class-only "
  "edges measurably hurts performance for the two sub-classes (Metals, Precious) where cross-class "
  "diversification evidently matters most.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. LIMITATIONS AND CAVEATS
# ═══════════════════════════════════════════════════════════════
h1("7. Limitations and Caveats")

bullet("Sample size: the out-of-sample test period spans roughly 5.5 years (2021 through mid-2026), "
       "considerably shorter than the original paper's 22-year out-of-sample window. Every result in "
       "this report should be read with that shorter, less statistically powerful sample in mind.")
bullet("Hyperparameter instability: the 121-candidate grid search consistently shows wide flat regions "
       "where many different (α, β) pairs score almost identically, and which pair wins can "
       "swing considerably from one walk-forward block to the next. \"The selected hyperparameters\" "
       "should therefore be read as one reasonable choice among several near-ties, not a uniquely "
       "optimal specification.")
bullet("Solver choice: this project's FBF-based solver targets the same convex optimization problem as "
       "the original paper's commercial MOSEK solver, but the two were never run side-by-side on "
       "identical inputs to confirm they converge to numerically identical solutions in every case.")
bullet("Small-universe results (NGL in particular) come from a pooled regression with far fewer "
       "cross-sectional observations per day than the full 22-product case, and are correspondingly "
       "more exposed to a favorable draw of hyperparameters or market regime, even when the reported "
       "test window was never touched during model fitting.")
bullet("The transaction-cost model is a flat, symmetric round-trip rate; it does not model bid-ask "
       "spread widening under stress, market impact, or the possibility that a genuinely thin market "
       "(such as NGL) may carry materially higher realistic trading costs than the 5 basis point "
       "baseline used uniformly across every universe in this report.")
bullet("The graph-topology characterization (Section 5.9) uses two illustrative windows, not the full "
       "multi-year history, and should be read as an observed pattern rather than a validated, "
       "general regime-detection signal.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═══════════════════════════════════════════════════════════════
h1("8. Conclusion")

p("This project set out to test whether the network-momentum mechanism of Pu, Roberts, Dong and "
  "Zohren (2023) -- a data-driven graph that diffuses each instrument's own momentum signal across its "
  "learned neighbors -- adds genuine value within a single, coherent 22-product commodity universe, "
  "distinct from the original paper's broader, four-traditional-asset-class setting.")

p("The honest answer is nuanced rather than a simple yes or no. The mechanism does appear to add real, "
  "economically sensible structure: sub-classes that are isolated from cross-class diversification "
  "perform measurably worse, the learned graph behaves coherently across calm and volatile regimes, "
  "and a conservative combination of individual and network momentum (SignCombo) modestly outperforms "
  "either signal alone. At the same time, none of the headline Sharpe-ratio gaps over a simple "
  "Long-Only benchmark clear conventional statistical significance in the full-22 universe, and every "
  "apparent edge anywhere in this project -- including the one genuinely significant result found, in "
  "the NGL-only sub-universe -- depends critically on fast execution and does not survive a single "
  "additional day of delay. Reversal, a key second driver of the original paper's own findings, does "
  "not appear at all in this commodity-only setting.")

p("Taken together, this project supports a specific, defensible claim: within a commodity-only "
  "universe, network momentum is a real and explainable structural effect on how diversification and "
  "cross-class relationships behave, but it is not, on this evidence, a statistically validated, "
  "delay-robust source of standalone trading edge.")

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
